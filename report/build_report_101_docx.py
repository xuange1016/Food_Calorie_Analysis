from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
RUN_DIR = PROJECT_ROOT / "training_runs" / "food101_101class"
METRICS_PATH = RUN_DIR / "evaluation_metrics.json"
HISTORY_PATH = RUN_DIR / "training_history.json"
SUMMARY_PATH = PROJECT_ROOT / "data" / "food101_101class" / "dataset_summary.json"
OUT_PATH = PROJECT_ROOT / "report" / "基于机器视觉的食堂外卖菜品识别与热量估算助手_101类模型报告.docx"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(45, 45, 45)
GRAY = RGBColor(90, 90, 90)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F3F5F7"


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def set_run_font(run, name="宋体", size=11, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_width(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for margin_name, value in [("top", "90"), ("bottom", "90"), ("start", "120"), ("end", "120")]:
                node = tc_mar.find(qn(f"w:{margin_name}"))
                if node is None:
                    node = OxmlElement(f"w:{margin_name}")
                    tc_mar.append(node)
                node.set(qn("w:w"), value)
                node.set(qn("w:type"), "dxa")


def add_heading(doc, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
    else:
        set_run_font(run, size=13, bold=True, color=BLUE)


def add_body(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=11, color=DARK)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=11)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], LIGHT_BLUE)
        set_cell_text(hdr[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 20 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, align=align)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("人工智能概论课程设计报告")
    set_run_font(r, size=18, bold=True, color=BLUE)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于机器视觉的食堂/外卖菜品识别与热量估算助手")
    set_run_font(r, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Food-101 全量 101 类模型版本")
    set_run_font(r, size=14, bold=True, color=BLUE)

    for _ in range(5):
        doc.add_paragraph()

    for label in ["小组成员：____________________________", "学号：____________________________", "班级：____________________________", "提交日期：2026 年 5 月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, size=12)

    doc.add_page_break()


def format_pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def main() -> None:
    metrics = load_json(METRICS_PATH)
    history = load_json(HISTORY_PATH)
    summary = load_json(SUMMARY_PATH)

    best_epoch = max(history, key=lambda item: item["val_acc"])
    final_epoch = history[-1]
    train_count = sum(item["train"] for item in summary.values())
    val_count = sum(item["val"] for item in summary.values())
    test_count = sum(item["test"] for item in summary.values())

    per_class = metrics["per_class"]
    sorted_classes = sorted(per_class.items(), key=lambda item: item[1]["accuracy"], reverse=True)
    top10 = sorted_classes[:10]
    bottom10 = sorted_classes[-10:]

    doc = Document()
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    add_cover(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Food-101 101类菜品识别与热量估算助手")
    set_run_font(run, size=9, color=GRAY)

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "本报告在原 10 类菜品识别系统基础上，将模型扩展到 Food-101 数据集完整 101 个食物类别。"
        "系统仍采用 ResNet18 迁移学习方案，结合 Flask 后端和 HTML/CSS/JavaScript 前端，实现图片上传、"
        "101 类食物识别、中文类别展示、营养信息估算和饮食建议输出。本版本完成 60600 张训练图片、15150 张验证图片、"
        "25250 张测试图片的全量训练与评估，测试集准确率达到 66.66%。",
    )
    add_body(doc, "关键词：Food-101；101类食物识别；ResNet18；迁移学习；Flask；营养估算")

    add_heading(doc, "一、项目升级背景", 1)
    add_body(
        doc,
        "原系统使用 Food-101 中筛选出的 10 类食物完成课程设计演示，能够验证从数据处理、模型训练、后端接口到前端展示的完整流程。"
        "为了进一步提升系统覆盖范围，本次升级将训练类别扩展为 Food-101 全量 101 类，并补充中文映射、营养信息和饮食建议配置，"
        "使系统更接近真实食物识别应用场景。",
    )

    add_heading(doc, "二、数据集与切分", 1)
    add_body(
        doc,
        "本次训练使用 Food-101 完整数据集。Food-101 共包含 101 个食物类别，每类约 1000 张图片。"
        "项目沿用官方测试集划分，并从官方训练集中按固定随机种子切分 20% 作为验证集。",
    )
    add_table(
        doc,
        ["项目", "数量"],
        [["类别数", "101"], ["训练集", str(train_count)], ["验证集", str(val_count)], ["测试集", str(test_count)]],
        widths=[2.4, 2.6],
    )
    add_bullet(doc, "数据整理脚本：training/dataset_prepare.py，新增 --all-classes 参数。")
    add_bullet(doc, "类别顺序文件：training_runs/food101_101class/class_names_101.json。")
    add_bullet(doc, "数据目录：data/food101_101class/train、val、test。")

    add_heading(doc, "三、模型训练方案", 1)
    add_body(
        doc,
        "模型继续采用 ResNet18，并加载 ImageNet 预训练权重进行迁移学习。训练脚本新增每轮 checkpoint 保存、"
        "checkpoint_latest 自动续训、checkpoint_best 最佳模型保存、训练历史 JSON 输出和 AMP 混合精度支持，"
        "以降低长时间训练中断造成的损失。",
    )
    add_table(
        doc,
        ["配置项", "取值"],
        [
            ["模型", "ResNet18"],
            ["类别数", "101"],
            ["图像尺寸", str(metrics["image_size"])],
            ["Epoch", str(final_epoch["epoch"])],
            ["Batch size", "32"],
            ["学习率", "0.001"],
            ["设备", "NVIDIA CUDA"],
            ["最佳 epoch", str(best_epoch["epoch"])],
        ],
        widths=[2.6, 3.2],
    )

    add_heading(doc, "四、训练与评估结果", 1)
    add_body(
        doc,
        "训练完成后，最终模型保存验证集准确率最高的权重。101 类任务难度显著高于原 10 类任务，"
        "模型在完整 Food-101 测试集上取得 66.66% 准确率，说明系统已经具备较宽类别覆盖的基础识别能力。",
    )
    add_table(
        doc,
        ["指标", "结果"],
        [
            ["最佳验证准确率", format_pct(best_epoch["val_acc"])],
            ["最佳 epoch", str(best_epoch["epoch"])],
            ["最终训练准确率", format_pct(final_epoch["train_acc"])],
            ["测试集准确率", format_pct(metrics["test_accuracy"])],
            ["测试样本数", str(metrics["test_samples"])],
            ["模型文件", "training_runs/food101_101class/food_model_101class.pth"],
        ],
        widths=[2.8, 3.8],
    )

    add_heading(doc, "五、类别表现分析", 1)
    add_body(doc, "从测试结果看，部分外观特征清晰的类别识别效果较好，例如毛豆、洋葱圈、牡蛎、马卡龙等；部分外观差异较大或与其他类别相似的类别表现较弱，例如苹果派、巧克力慕斯、猪排、早餐卷饼等。")
    add_table(
        doc,
        ["表现较好类别", "准确率", "正确/总数"],
        [[name, format_pct(item["accuracy"]), f"{item['correct']}/{item['total']}"] for name, item in top10],
        widths=[2.8, 1.4, 1.4],
    )
    add_table(
        doc,
        ["表现较弱类别", "准确率", "正确/总数"],
        [[name, format_pct(item["accuracy"]), f"{item['correct']}/{item['total']}"] for name, item in bottom10],
        widths=[2.8, 1.4, 1.4],
    )

    add_heading(doc, "六、系统功能升级", 1)
    add_bullet(doc, "后端默认优先加载 101 类模型，模型不存在时回退到旧 10 类模型。")
    add_bullet(doc, "新增 backend/food_display_names.json，提供 101 类英文标签到中文名称的映射。")
    add_bullet(doc, "完善 backend/food_nutrition.json，为 101 类提供营养估算、食物分类和饮食建议。")
    add_bullet(doc, "Web 前端显示当前模型类别数、输入尺寸和运行设备。")
    add_bullet(doc, "预测结果和 Top-5 候选均显示中文名称，提升演示可读性。")

    add_heading(doc, "七、系统局限与改进方向", 1)
    add_body(
        doc,
        "本系统当前仍属于课程设计演示版本。营养信息使用每 100 克平均估算值，无法根据图片自动判断实际重量；"
        "模型采用单标签分类方式，无法处理一张图中同时出现多种菜品；ResNet18 在 101 类任务上存在一定泛化上限。"
        "后续可考虑引入更强的 EfficientNet、ConvNeXt 或 ViT 模型，增加多菜品检测和分割能力，并建立更精确的营养数据库。",
    )

    add_heading(doc, "八、结论", 1)
    add_body(
        doc,
        "本次升级完成了 Food-101 全量 101 类模型训练、评估和 Web 服务集成。系统能够识别更丰富的食物类别，"
        "并以中文形式展示识别结果、营养信息和饮食建议。与原 10 类版本相比，101 类版本覆盖范围大幅提升，"
        "更适合展示完整机器视觉应用从数据集扩展、训练脚本增强到 Web 系统落地的全过程。",
    )

    add_heading(doc, "参考资料", 1)
    add_body(doc, "[1] Food-101 数据集：https://data.vision.ee.ethz.ch/cvl/datasets_extra/food-101/")
    add_body(doc, "[2] PyTorch 官方文档：https://pytorch.org/")
    add_body(doc, "[3] Torchvision Models 文档：https://pytorch.org/vision/stable/models.html")
    add_body(doc, "[4] Flask 官方文档：https://flask.palletsprojects.com/")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
