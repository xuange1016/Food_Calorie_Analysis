from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = PROJECT_ROOT / "report" / "基于机器视觉的食堂外卖菜品识别与热量估算助手_课程设计报告.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"


def set_run_font(run, name="宋体", size=11, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_width(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for margin_name, value in [("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")]:
                node = tc_mar.find(qn(f"w:{margin_name}"))
                if node is None:
                    node = OxmlElement(f"w:{margin_name}")
                    tc_mar.append(node)
                node.set(qn("w:w"), value)
                node.set(qn("w:type"), "dxa")


def add_heading(doc, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)


def add_body(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=11)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], LIGHT_GRAY)
        set_cell_text(hdr[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(value) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, align=align)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("人工智能概论课程")
    set_run_font(r, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Introduction to Artificial Intelligence")
    set_run_font(r, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KCBZ-0013100403")
    set_run_font(r, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2025-2026（2）")
    set_run_font(r, size=10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("课程设计报告")
    set_run_font(r, size=16, bold=True)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目名称：基于机器视觉的食堂/外卖菜品识别与热量估算助手")
    set_run_font(r, size=14, bold=True)

    doc.add_paragraph()
    for label in ["小组成员：____________________________", "学号：____________________________", "班级：____________________________", "提交日期：2026 年 5 月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, size=12)

    doc.add_page_break()


def add_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("基于机器视觉的食堂/外卖菜品识别与热量估算助手")
    set_run_font(run, size=9, color=GRAY)


def build_doc() -> None:
    doc = Document()
    add_cover(doc)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    add_heading(doc, "摘要", 1)
    add_body(doc, "本项目面向大学生食堂和外卖饮食场景，设计并实现了一个基于机器视觉的菜品识别与热量估算助手。系统使用 Food-101 数据集中的 10 类食物图片训练 ResNet18 图像分类模型，用户上传食物图片后，后端调用模型识别食物类别，并根据本地营养信息表估算热量、蛋白质、脂肪和碳水化合物，最后生成简单饮食建议。项目最终实现了一个可运行的 Web demo，测试集准确率达到 84.04%，能够满足人工智能概论课程设计中机器视觉应用、数据处理、模型训练、系统实现和成果展示的要求。")
    add_body(doc, "关键词：机器视觉；食物识别；迁移学习；ResNet18；Flask；热量估算")

    add_heading(doc, "一、项目背景", 1)
    add_body(doc, "随着外卖和校园食堂场景的普及，大学生每天都会面对多种饮食选择。许多学生希望了解一餐的大致热量和营养结构，但手动查询每种食物的营养信息比较麻烦。传统饮食记录软件通常依赖用户手动输入食物名称，使用门槛较高。")
    add_body(doc, "本项目选择机器视觉应用方向，设计并实现一个食物图片识别与营养估算 Web demo。用户上传一张菜品图片后，系统自动识别食物类别，并结合本地平均营养信息表给出热量和营养估算。项目重点体现人工智能应用的完整流程，而不是追求医学级精准营养计算。")

    add_heading(doc, "二、数据集介绍", 1)
    add_body(doc, "本项目使用公开食物图像分类数据集 Food-101。Food-101 原始数据集包含 101 个食物类别，每个类别约 1000 张图片，常用于食物图像分类研究和教学实践。为了控制课程设计难度并保证训练效率，本项目第一版从 Food-101 中筛选 10 类常见食物进行训练和演示。")
    add_table(doc, ["英文类别", "中文含义"], [
        ["caesar_salad", "凯撒沙拉"],
        ["chicken_wings", "鸡翅/炸鸡类"],
        ["chocolate_cake", "巧克力蛋糕"],
        ["dumplings", "饺子"],
        ["filet_mignon", "牛排"],
        ["fried_rice", "炒饭"],
        ["hamburger", "汉堡"],
        ["pizza", "披萨"],
        ["ramen", "拉面"],
        ["sushi", "寿司"],
    ], widths=[2.8, 3.4])
    add_table(doc, ["数据集划分", "图片数量"], [["训练集", "6000"], ["验证集", "1500"], ["测试集", "2500"]], widths=[3.2, 3.0])

    add_heading(doc, "三、数据预处理", 1)
    add_body(doc, "Food-101 原始数据中包含 images 和 meta 目录。meta/train.txt 与 meta/test.txt 提供官方训练集和测试集划分信息。项目通过 training/dataset_prepare.py 读取这些元数据，根据 selected_classes.json 筛选目标类别，并将训练数据进一步划分出验证集。")
    add_bullet(doc, "读取 Food-101 原始 train.txt 和 test.txt。")
    add_bullet(doc, "筛选 10 个目标类别。")
    add_bullet(doc, "从原始训练集中划分 20% 作为验证集。")
    add_bullet(doc, "将图片整理为 train、val、test 三个目录。")
    add_bullet(doc, "保存类别顺序到 backend/models/class_names.json，保证训练和预测一致。")
    add_body(doc, "训练阶段采用随机水平翻转等轻量数据增强，并将图片缩放到 224×224 后进行 ImageNet 均值和标准差归一化。验证集和测试集不使用随机增强，以保证评估结果稳定。")

    add_heading(doc, "四、模型设计", 1)
    add_body(doc, "本项目采用 PyTorch 框架实现食物图像分类模型，模型选择 ResNet18，并使用迁移学习方法。ResNet18 是经典卷积神经网络结构，通过残差连接缓解深层网络训练困难问题。项目加载在 ImageNet 上训练得到的预训练权重，将最后的全连接分类层替换为 10 类输出，再使用筛选后的 Food-101 子集进行微调训练。")
    add_table(doc, ["参数", "设置"], [
        ["模型", "ResNet18"],
        ["图像尺寸", "224×224"],
        ["类别数", "10"],
        ["Epoch", "8"],
        ["Batch size", "32"],
        ["优化器", "Adam"],
        ["学习率", "0.001"],
        ["设备", "NVIDIA GPU / CUDA"],
    ], widths=[2.4, 3.8])
    add_body(doc, "训练过程中记录训练集损失、训练集准确率、验证集损失和验证集准确率，并保存验证集准确率最高的模型。最终模型文件保存为 backend/models/food_model.pth。")

    add_heading(doc, "五、系统功能设计", 1)
    add_body(doc, "系统分为前端展示模块、Flask 后端接口模块、模型预测模块、营养信息查询与建议模块。用户上传图片后，前端提交到后端；后端保存临时图片并调用模型预测；模型返回类别和置信度；系统根据类别查询本地营养信息表并返回 JSON；前端展示识别结果、营养信息和饮食建议。")
    add_table(doc, ["模块", "主要功能"], [
        ["前端模块", "图片上传、预览、调用接口、展示识别和营养结果"],
        ["后端接口", "接收图片、异常处理、组织 JSON 返回"],
        ["模型预测", "加载 ResNet18 模型并输出类别、置信度和 Top-3"],
        ["营养模块", "根据类别查询热量、蛋白质、脂肪、碳水和建议"],
    ], widths=[2.1, 4.1])

    add_heading(doc, "六、系统实现", 1)
    add_body(doc, "后端使用 Flask 实现图片上传接口 POST /predict。接口会检查是否上传图片、图片格式是否正确、模型文件是否存在以及营养信息是否匹配。预测逻辑封装在 backend/predict.py 中，营养查询逻辑封装在 backend/nutrition.py 中。")
    add_body(doc, "前端使用 HTML、CSS 和 JavaScript 实现。页面包含上传区域、图片预览区、识别结果卡片、营养信息卡片、饮食建议和 Top-3 预测结果。为了兼容手机浏览器，图片上传控件使用 accept=\"image/*\" 和 capture=\"environment\"，部分手机浏览器可直接调用相机。")
    add_body(doc, "项目目录结构如下：")
    add_table(doc, ["目录/文件", "作用"], [
        ["backend/", "Flask 后端、预测模块、营养模块和模型文件"],
        ["frontend/", "Web 页面、样式和交互脚本"],
        ["training/", "数据整理、训练、评估和单图预测脚本"],
        ["data/food101_10class/", "筛选后的 10 类训练、验证和测试数据"],
        ["report/", "课程设计报告文件"],
    ], widths=[2.6, 3.6])

    add_heading(doc, "七、项目创新点", 1)
    add_bullet(doc, "选题贴近大学生日常食堂和外卖饮食场景，比常见分类项目更有应用价值。")
    add_bullet(doc, "将图像分类结果与营养信息表结合，实现从识别到建议的应用闭环。")
    add_bullet(doc, "采用迁移学习降低训练成本，适合课程设计环境。")
    add_bullet(doc, "提供 Web demo，可现场上传图片展示识别和营养估算结果。")
    add_bullet(doc, "前端保留手机浏览器拍照上传能力，为后续打包手机应用留下扩展空间。")

    add_heading(doc, "八、项目成果展示", 1)
    add_body(doc, "项目最终实现了一个可运行的 Web demo。用户可以在浏览器中上传食物图片，系统返回识别类别、置信度、Top-3 预测、热量、蛋白质、脂肪、碳水化合物和饮食建议。")
    add_table(doc, ["指标", "结果"], [
        ["最佳验证准确率", "78.60%"],
        ["测试集准确率", "84.04%"],
        ["测试样本数", "2500"],
        ["模型文件", "backend/models/food_model.pth"],
    ], widths=[2.8, 3.4])
    add_table(doc, ["类别", "测试准确率"], [
        ["caesar_salad", "78.80%"],
        ["chicken_wings", "91.20%"],
        ["chocolate_cake", "85.20%"],
        ["dumplings", "89.20%"],
        ["filet_mignon", "79.60%"],
        ["fried_rice", "92.80%"],
        ["hamburger", "87.60%"],
        ["pizza", "94.80%"],
        ["ramen", "70.00%"],
        ["sushi", "71.20%"],
    ], widths=[3.0, 3.0])
    add_body(doc, "示例结果：上传披萨图片时，系统预测类别为 pizza，置信度为 100.00%，并返回每 100 克估算热量 266 kcal、蛋白质 11 g、脂肪 10 g、碳水 33 g，建议为“披萨热量较高，建议搭配沙拉并控制摄入量”。")
    add_body(doc, "项目展示材料包括 Web 首页、上传图片预览、识别结果、训练过程和测试准确率等截图。")

    add_heading(doc, "九、项目分析总结", 1)
    add_body(doc, "本项目完成了一个基于机器视觉的菜品识别与热量估算系统。项目通过 Food-101 数据集训练 ResNet18 食物分类模型，并将模型集成到 Flask 后端和 Web 前端中，实现了从图片上传、模型预测、营养查询到结果展示的完整流程。")
    add_body(doc, "从实验结果看，模型在 10 类 Food-101 子集上的测试准确率达到 84.04%，说明迁移学习方法能够在较短训练时间内取得较好的分类效果。不同类别之间准确率存在差异，其中 pizza、fried_rice、chicken_wings 等类别表现较好，而 ramen 和 sushi 准确率相对较低，可能与图像背景复杂、类别间视觉相似度较高、样本多样性较强有关。")
    add_body(doc, "项目仍存在一些不足：当前只能识别 10 类食物，无法估算食物重量，无法处理一张图片中的多个菜品，营养信息来自平均值表，不具备医学级精度。后续可以扩展到更多类别，引入目标检测或图像分割，增加手动修正功能，并尝试将模型转换为移动端推理格式。")
    add_body(doc, "总体而言，本项目具有明确的应用场景、完整的 AI 流程和可运行的系统 demo，能够满足《人工智能概论》课程设计中关于数据集、项目报告、系统源代码和项目 demo 的提交要求。")

    add_heading(doc, "参考资料", 1)
    add_body(doc, "[1] Food-101 数据集：https://data.vision.ee.ethz.ch/cvl/datasets_extra/food-101/")
    add_body(doc, "[2] PyTorch 官方文档：https://pytorch.org/")
    add_body(doc, "[3] Torchvision Models 文档：https://pytorch.org/vision/stable/models.html")
    add_body(doc, "[4] Flask 官方文档：https://flask.palletsprojects.com/")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_doc()
